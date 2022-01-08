from tkinter import *
import warnings
from PIL import ImageTk,Image
from PyPDF2 import pdf
from Frame import login
from tkinter import filedialog
from tkinter.messagebox import showinfo, WARNING
from tkinter.scrolledtext import ScrolledText
import PyPDF2,os,requests,openpyxl,time,json,csv
from selenium import webdriver
from selenium.webdriver.common.keys import Keys

class Mainhome:
    def __init__(self):
        w=Tk()
        w.geometry('1400x700') 
        w.configure(bg='#000000')
        w.resizable(False, False) 
        w.title("Filo")
        w.iconbitmap(r'Frame/home_img/icon.ico')
        def default_home():
            homeframe=Frame(w,width=1400,height=658,bg='#262626')
            glabel=Label(w,text='Home',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
            vector_img = PhotoImage(file = f"Frame/home_img/Vector.png")
            vector = Label(image=vector_img,background='#262626')
            vector.image=vector_img
            vector.place(x=180,y=160)
            Label(homeframe,text='Welcome to Filo!',font=("Poppins",80),background='#262626',foreground='#FFFFFF').place(x=350,y=100)
            Label(homeframe,text='Manage Docs, PDF, CSV files and much more from one place!',font=("Poppins",30),background='#262626',foreground='#FFFFFF').place(x=70,y=300)   
            homeframe.place(x=0,y=42)

        def toggle_win():
            global menu
            menu=Frame(w,width=300,height=700,bg='#D2E6FB')
            menu.place(x=0,y=0)

            def bttn(x,y,text,bcolor,fcolor,cmd):
                def on_entera(e):
                    myButton1['background']=bcolor
                    myButton1['foreground']='#262626'
                
                def on_leavea(e):
                    myButton1['background']=fcolor
                    myButton1['foreground']='#262626'

                myButton1= Button(menu,text=text,width=25,height=1,fg='#262626',border=0,font=("Poppins",15),bg=fcolor,activebackground='#262626',activeforeground=bcolor,command=cmd)
                myButton1.bind("<Enter>",on_entera)
                myButton1.bind("<Enter>",on_leavea)
                myButton1.place(x=x,y=y)

            def pdf_code():
                #function to merge pdfs
                def pdfmerge():
                    mergeframe=Frame(pdfframe,width=1400,height=658,bg='#000000')
                    mergeframe.place(x=700,y=0)
                    Label(mergeframe,text='PDF Merge',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=250,y=0)
                    global pdfmergecount
                    pdfmergecount=0
                    def selectfile():
                        global pdfmergecount,mergefile1,mergefile2
                        if(pdfmergecount!=2):
                            if pdfmergecount==0:
                                mergefile1=filedialog.askopenfilename(initialdir="C:\\",title="Open",filetypes=(("PDF Files","*.pdf"),("All files","*.*")))
                                if mergefile1!='' and mergefile1:
                                    pdfmergecount+=1
                                    t1=Entry(mergeframe,bd=0,highlightthickness=0,font=("Poppins",16))
                                    t1.place(x=100,y=230,width=650)
                                    t1.insert(0,os.path.basename(mergefile1))
                                    t1.configure(state='readonly',readonlybackground="#000000",foreground="#FFFFFF")
                            else:
                                mergefile2=filedialog.askopenfilename(initialdir="C:\\",title="Open",filetypes=(("PDF Files","*.pdf"),("All files","*.*")))
                                if mergefile2!='':
                                    pdfmergecount+=1
                                    t1=Entry(mergeframe,bd=0,highlightthickness=0,font=("Poppins",16))
                                    t1.place(x=100,y=330,width=650)
                                    t1.insert(0,os.path.basename(mergefile2))
                                    t1.configure(state='readonly',readonlybackground="#000000",foreground="#FFFFFF")
                        else:
                            showinfo(title='Error',message='Maximum two files Supported.',icon=WARNING)
                    def finalmerge():
                        if pdfmergecount!=2:
                            showinfo(title='Error',message='Merge requires two files.',icon=WARNING)
                        else:
                            pdf1File = open(mergefile1, 'rb')
                            pdf2File = open(mergefile2, 'rb')
                            # File Readers
                            pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
                            pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
                            #File writer
                            if pdf1Reader.isEncrypted==False and pdf2Reader.isEncrypted==False:
                                pdfWriter = PyPDF2.PdfFileWriter()
                                for pageNum in range(pdf1Reader.numPages):
                                    pageObj = pdf1Reader.getPage(pageNum)
                                    pdfWriter.addPage(pageObj)
                                
                                for pageNum in range(pdf2Reader.numPages):
                                    pageObj = pdf2Reader.getPage(pageNum)
                                    pdfWriter.addPage(pageObj)
                                pdfPath = filedialog.asksaveasfilename(initialdir="C:\\",defaultextension = "*.pdf", filetypes = (("PDF Files", "*.pdf"),("All files","*.*")))
                                if pdfPath: #If the user only selects the save file location
                                    pdfOutputFile = open(pdfPath, 'wb')
                                    pdfWriter.write(pdfOutputFile)
                                    pdfOutputFile.close()
                                    pdf1File.close()
                                    pdf2File.close()
                                    showinfo(title='Success', message='File Merged Successfully!', icon='info')
                                    pdfmerge()
                            else:
                                showinfo(title='Error', message='Encrypted Files Not Supported.', icon=WARNING)
                                pdfmerge()                           

                    Button(mergeframe,text="Merge & Save",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = finalmerge,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=450,y=558,width=200)
                    Button(mergeframe,text="Select File",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = selectfile,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=100,y=100,width=500)

                #funtion to split pdfs
                def pdfsplit():
                    splitframe=Frame(pdfframe,width=1400,height=658,bg='#000000')
                    splitframe.place(x=700,y=0)
                    Label(splitframe,text='Split PDF',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=250,y=0)
                    Label(splitframe,text='From',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=40,y=240)
                    Label(splitframe,text='To',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=370,y=240)
                    fromsplit=Entry(splitframe,font=("Poppins",18),background="#FFFFFF")
                    fromsplit.place(x=140,y=250,width=80,height=40)
                    tosplit=Entry(splitframe,font=("Poppins",18),background="#FFFFFF")
                    tosplit.place(x=440,y=250,width=80,height=40)
                    def selectsplitfile():
                        global splitfile,splitcount
                        splitcount=0
                        splitfile=filedialog.askopenfilename(initialdir="C:\\",title="Open",filetypes=(("PDF Files","*.pdf"),("All files","*.*")))
                        if splitfile:
                            splitcount+=1
                            t1=Entry(splitframe,bd=0,highlightthickness=0,font=("Poppins",16))
                            t1.place(x=100,y=330,width=650)
                            t1.insert(0,os.path.basename(splitfile))
                            t1.configure(state='readonly',readonlybackground="#000000",foreground="#FFFFFF")
                            
                    def splitprocess():
                        if splitcount!=0:
                            pdfFile = open(splitfile,'rb')
                            pdfReader = PyPDF2.PdfFileReader(pdfFile)
                            pdfWriter = PyPDF2.PdfFileWriter()
                            pdfrange=pdfReader.numPages
                            pageSt=fromsplit.get()
                            pageEnd=tosplit.get()
                            if pageSt!='' and pageEnd!='':
                                if pageSt.isdigit() and pageEnd.isdigit():
                                    if int(pageSt)<=pdfrange and int(pageEnd)<=pdfrange:
                                        if int(pageSt)<=int(pageEnd):
                                            for i in range(int(pageSt)-1,int(pageEnd)):
                                                pageObj = pdfReader.getPage(i)
                                                pdfWriter.addPage(pageObj)
                                            
                                            splitwrite = filedialog.asksaveasfilename(initialdir="C:\\",defaultextension = "*.pdf", filetypes = (("PDF Files", "*.pdf"),("All files","*.*")))
                                            if splitwrite:
                                                pdfOutputFile = open(splitwrite, 'wb')
                                                pdfWriter.write(pdfOutputFile)
                                                pdfOutputFile.close()
                                                pdfFile.close()
                                                showinfo(title='Success', message='File has been split Successfully!', icon='info')
                                                pdfsplit()
                                        else:
                                            showinfo(title='Error', message='Invalid Range', icon=WARNING)
                                    else:
                                        showinfo(title='Error', message='Invalid Range', icon=WARNING)
                                else:
                                    showinfo(title='Error', message='Invalid Range', icon=WARNING)
                                    fromsplit.delete(0, END)
                                    tosplit.delete(0, END)
                            else:
                                showinfo(title='Error', message='Range cannot be empty', icon=WARNING)
                        else:
                            showinfo(title='Error', message='Select a file to Split', icon=WARNING)


                    Button(splitframe,text="Split & Save",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = splitprocess,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=450,y=558,width=200)
                    Button(splitframe,text="Select File",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = selectsplitfile,background="#FFFFFF",foreground="#262626",activebackground="#FFFFFF",relief = "flat").place(x=100,y=100,width=500)

                menu.destroy()
                glabel=Label(w,text='PDF Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                pdfframe=Frame(w,width=1400,height=658,bg='#262626')
                pdfframe.place(x=0,y=42)
                #merge pdf button
                merge_img = PhotoImage(file = f"Frame/home_img/merge.png")
                label = Label(image=merge_img)
                label.image=merge_img
                mergebtn = Button(pdfframe,image = merge_img,borderwidth = 0,highlightthickness = 0,command = pdfmerge,background="#262626",activebackground="#262626",relief = "flat").place(x=130,y=240,width=129,height=119)
                
                #split pdf button
                split_img = PhotoImage(file = f"Frame/home_img/split.png")
                label = Label(image=split_img)
                label.image=split_img
                splitbtn = Button(pdfframe,image = split_img,borderwidth = 0,highlightthickness = 0,command = pdfsplit,background="#262626",activebackground="#262626",relief = "flat").place(x=330,y=240,width=129,height=119)
                pdfmerge()

            def word_code():
                menu.destroy()
                def textodoc():
                    doctext=textentry.get("1.0",END)
                    import docx,pyperclip
                    doc=docx.Document()
                    header=headerentry.get()
                    if header!='':
                        doc.add_heading(header,0)
                    text=doctext.split('\n')
                    for i in range(len(text)):
                        doc.add_paragraph(text[i])
                    wordPath = filedialog.asksaveasfilename(defaultextension = "*.docx", filetypes = (("Word Files", "*.docx"),))
                    print(wordPath)
                    if wordPath:
                        doc.save(wordPath)
                    showinfo(title='Success', message='File created Successfully!', icon='info')
                    headerentry.delete(0, END)
                    textentry.delete("1.0",END)
                    word_code()
                    
                wordframe=Frame(w,width=1400,height=658,bg='#262626')
                glabel=Label(w,text='Word Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
                wordframe.place(x=0,y=42)
                #entry of header
                Label(wordframe,text='Header',font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=600,y=0)
                headerentry=Entry(wordframe,font=("Poppins",20),background="#FFFFFF")
                headerentry.place(x=600,y=50,width=790,height=60)
                #content entry
                textentry=ScrolledText(wordframe,font=("Poppins",20),background="#FFFFFF")
                textentry.place(x=600,y=120,width=790,height=508)

                # Label(wordframe,text='Font Size',font=("Poppins",18),background='#262626',foreground='#FFFFFF').place(x=335,y=160)
                # fontsize=Entry(wordframe,font=("Poppins",18),background="#FFFFFF")
                # fontsize.place(x=340,y=200,width=80,height=40)
                
                # Label(wordframe,text='Font Style',font=("Poppins",18),background='#262626',foreground='#FFFFFF').place(x=100,y=160)
                # fontstyleval = StringVar()
                # fontstyle = ttk.Combobox(wordframe,textvariable = fontstyleval,foreground="#000000",font=("Poppins",10),state='readonly')
                # fontstyle.place(x=100,y=200,height=40)
                # fontstyle['values'] = ('Cambria','Comic Sans MS','Calibri','Cavolini','Arial')
                # fontstyle.current(1)

                doc_img = PhotoImage(file = f"Frame/home_img/doc.png")
                label = Label(image=doc_img)
                label.image=doc_img
                Button(wordframe,image=doc_img,borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = textodoc,background="#262626",foreground="#000000",activebackground="#262626",relief = "flat").place(x=230,y=300,width=129,height=119)

            def home():
                menu.destroy()
                default_home()

            def excel_code():
                menu.destroy()
                excelframe=Frame(w,width=1400,height=658,bg='#262626')
                excelframe.place(x=0,y=42)

                cityentry=Entry(excelframe,font=("Poppins",20),background="#FFFFFF")
                cityentry.place(x=450,y=80,width=400,height=40)
                dataframe=Frame(excelframe,width=1400,height=450,bg='#262626')
                dataframe.place(x=0,y=135)
                def labeldata(city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed):
                    Label(dataframe,background="#262626").place(x=0,y=0,width=1400,height=450)
                    Label(dataframe,text='City: '+str(city_name),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=0)
                    Label(dataframe,text='Latitude: '+str(latitude),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=50)
                    Label(dataframe,text='Longitude: '+str(longitude),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=100)
                    Label(dataframe,text='Temperature: '+str(temp),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=150)
                    Label(dataframe,text='Pressure: '+str(pressure),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=200)
                    Label(dataframe,text='Humidity: '+str(humidity),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=250)
                    Label(dataframe,text='Country: '+str(country),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=300)
                    Label(dataframe,text='Sky: '+str(desc),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=350)
                    Label(dataframe,text='Wind Speed: '+str(wind_speed),font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=10,y=400)                
                def get_weather():
                    #API to get weather details from openweathermap.org
                    api_key = "fc997859c662f09ea9dac60a3c71ecbc"
                    apiurl = "http://api.openweathermap.org/data/2.5/weather?"
                    city_name = cityentry.get()
                    if city_name=='':
                        showinfo(title='Error', message='Enter a city!', icon=WARNING)
                    else:   

                        url = apiurl + "appid=" + api_key + "&q=" + city_name
                        try:
                            response = requests.get(url)
                        except:
                            showinfo(title='Error', message='Internet Connection unavailable!', icon=WARNING)
                        else:    
                            x = json.loads(response.text)
                            if x["cod"] != "404":

                                y = x["main"]
                                z = x["weather"]
                                a = x['coord']
                                b = x['wind']
                                c = x['sys'] 

                                latitude=a['lat']
                                longitude=a['lon']
                                temp = y["temp"]
                                pressure = y["pressure"]
                                humidity = y["humidity"]
                                country=c['country']
                                desc = z[0]["description"]
                                wind_speed=b['speed']
                        
                                labeldata(city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed)
                            else:
                                showinfo(title='Error', message='City not found!', icon=WARNING)
                def export_details():
                    #API to get weather details from openweathermap.org
                    api_key = "fc997859c662f09ea9dac60a3c71ecbc"
                    apiurl = "http://api.openweathermap.org/data/2.5/weather?"
                    city_name = cityentry.get()
                    if city_name=='':
                        showinfo(title='Error', message='Enter a city!', icon=WARNING)
                    else:    
                        url = apiurl + "appid=" + api_key + "&q=" + city_name
                        try:
                            response = requests.get(url)
                        except:
                            showinfo(title='Error', message='Internet Connection unavailable!', icon=WARNING)
                        else:    
                            x = json.loads(response.text) # Converts json data to python dictionary

                            if x["cod"] != "404":          # Checks status_code of the response

                                y = x["main"]
                                z = x["weather"]
                                a = x['coord']
                                b = x['wind']
                                c = x['sys'] 

                                latitude=a['lat']
                                longitude=a['lon']
                                temp = y["temp"]
                                pressure = y["pressure"]
                                humidity = y["humidity"]
                                country=c['country']
                                desc = z[0]["description"]
                                wind_speed=b['speed']             
                                # Loads an excel file and adds the data to it row-wise
                                wb=openpyxl.load_workbook('FiloDirectory\\test.xlsx')
                                sheet = wb.get_active_sheet()
                                list=[city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed]
                                # Displays data in the window
                                labeldata(city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed)                 
                                j=sheet.get_highest_row()+1
                                sheet['A'+str(j)].value=j-1
                                for i in range(1,11):
                                    if i<=9:
                                        sheet.cell(row=j,column=i+1).value=list[i-1]
                                wb.save('FiloDirectory/test.xlsx')
                                showinfo(title='Successful', message='Data exported successfully!')
                                cityentry.delete(0,END)

                            else:
                                showinfo(title='Error', message='City not found!', icon=WARNING)
                def export_csv():
                    api_key = "fc997859c662f09ea9dac60a3c71ecbc"
                    apiurl = "http://api.openweathermap.org/data/2.5/weather?"
                    city_name = cityentry.get()
                    if city_name=='':
                        showinfo(title='Error', message='Enter a city!', icon=WARNING)
                    else:   

                        url = apiurl + "appid=" + api_key + "&q=" + city_name
                        try:    
                            response = requests.get(url)
                        except:
                            showinfo(title='Error', message='Internet Connection unavailable!', icon=WARNING)
                        else:
                            x = json.loads(response.text) # Converts json data to python dictionary
                            if x["cod"] != "404":        # Checks status_code of the response

                                y = x["main"]
                                z = x["weather"]
                                a = x['coord']
                                b = x['wind']
                                c = x['sys'] 

                                latitude=a['lat']
                                longitude=a['lon']
                                temp = y["temp"]
                                pressure = y["pressure"]
                                humidity = y["humidity"]
                                country=c['country']
                                desc = z[0]["description"]
                                wind_speed=b['speed']
                                list=[city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed]

                                labeldata(city_name,latitude,longitude,temp,pressure,humidity,country,desc,wind_speed)
                                # Loads csv file and appends the weather data to it row-wise
                                weatherfile = open('FiloDirectory/weather.csv', 'a', newline='')
                                weathWriter = csv.writer(weatherfile)
                                weathWriter.writerow(list)
                                weatherfile.close()
                                showinfo(title='Successful', message='Data exported successfully!')
                                cityentry.delete(0,END)
            

                search_img = PhotoImage(file = f"Frame/home_img/search.png")
                label = Label(image=search_img)
                label.image=search_img

                Label(excelframe,text='Place',font=("Poppins",24),background='#262626',foreground='#FFFFFF').place(x=450,y=10)
                Button(excelframe,borderwidth = 0,image=search_img,highlightthickness = 0,font=("Poppins",15),command = get_weather,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=852,y=80,width=40,height=40)
                
                Button(excelframe,borderwidth = 0,text="Export as Excel",highlightthickness = 0,font=("Poppins",15),command = export_details,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=510,y=600,height=40)
                Button(excelframe,borderwidth = 0,text="Export as CSV",highlightthickness = 0,font=("Poppins",15),command = export_csv,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=680,y=600,height=40)
                Label(w,text='Excel Manager',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)
               
            def services_code():
                def gmail():
                    def email_log():
                        search=email.get()
                        global driver
                        driver = webdriver.Chrome()
                        driver.get("https://gmail.com")
                        driver.find_element_by_id("identifierId").send_keys(search)
                    gmailframe=Frame(servframe,width=700,height=658,bg='#000000')
                    gmailframe.place(x=700,y=0)
                    gmaillabel_img = PhotoImage(file = f"Frame/home_img/gmaillabel.png")
                    label = Label(image=gmaillabel_img)
                    label.image=gmaillabel_img
                    Label(gmailframe,image=gmaillabel_img,font=("Poppins",24),background='#FFFFFF').place(x=0,y=0,width=700,height=80)

                    Label(gmailframe,text="Email",font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=0,y=170,width=700,height=80)
                    email = Entry(gmailframe,bd = 0,bg = "#FFFFFF",font=("Poppins",15),highlightthickness = 0)
                    email.place(x=180,y=230,width=350)

                    Button(gmailframe,text="Lets Go!",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = email_log,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=280,y=320,width=120)


                def amazon():
                    # Accessing the website amazon using selenium webdriver
                    def browse_amaz():
                        global driver
                        driver=webdriver.Chrome()
                        search=amazonprod.get()
                        driver.get('https://www.amazon.in')
                        time.sleep(1)
                        driver.find_element_by_id('twotabsearchtextbox').send_keys(search+Keys.ENTER)
                        time.sleep(3)
                        price =driver.find_element_by_class_name('a-price-whole').text
                        time.sleep(4)
                        driver.find_element_by_class_name('a-size-medium.a-color-base.a-text-normal').click()
                        Label(amazonframe,text="Price: "+str(price),font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=230,y=600)


                    amazonframe=Frame(servframe,width=700,height=658,bg='#000000')
                    amazonframe.place(x=700,y=0)
                    amazonlabel_img = PhotoImage(file = f"Frame/home_img/amazonlabel.png")
                    label = Label(image=amazonlabel_img)
                    label.image=amazonlabel_img
                    Label(amazonframe,image=amazonlabel_img,font=("Poppins",24),background='#FFFFFF').place(x=0,y=0,width=700,height=80)

                    Label(amazonframe,text="Search Product",font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=0,y=170,width=700,height=80)
                    
                    amazonprod = Entry(amazonframe,bd = 0,bg = "#FFFFFF",font=("Poppins",15),highlightthickness = 0)
                    amazonprod.place(x=180,y=230,width=350)

                    Button(amazonframe,text="Lets Go!",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = browse_amaz,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=280,y=320,width=120)

                def flipkart():
                    # Accessing the flipkart website using selenium webdriver

                    def flipk_browse():
                        global driver
                        search=flipkartprod.get()
                        driver=webdriver.Chrome()
                        driver.get('https://www.flipkart.com')
                        time.sleep(1)
                        driver.find_element_by_class_name('_2KpZ6l._2doB4z').click()
                        time.sleep(1)
                        driver.find_element_by_class_name('_3704LK').send_keys(search+Keys.ENTER)
                        time.sleep(4)
                        price =driver.find_element_by_class_name('_30jeq3._1_WHN1').text
                        Label(flipkartframe,text="Price: "+str(price),font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=230,y=600)        
                        
                    flipkartframe=Frame(servframe,width=700,height=658,bg='#000000')
                    flipkartframe.place(x=700,y=0)
                    flipkartlabel_img = PhotoImage(file = f"Frame/home_img/flipkartlabel.png")
                    label = Label(image=flipkartlabel_img)
                    label.image=flipkartlabel_img
                    Label(flipkartframe,image=flipkartlabel_img,font=("Poppins",24),background='#FFFFFF').place(x=0,y=0,width=700,height=80)

                    Label(flipkartframe,text="Search Product",font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=0,y=170,width=700,height=80)

                    flipkartprod = Entry(flipkartframe,bd = 0,bg = "#FFFFFF",font=("Poppins",15),highlightthickness = 0)
                    flipkartprod.place(x=180,y=230,width=350)
                    
                    Button(flipkartframe,text="Lets Go!",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = flipk_browse,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=280,y=320,width=120)

                def youtube():
                    # Accessing the flipkart website using selenium webdriver

                    def yt_browse():
                        search=videotxt.get()
                        global driver
                        driver=webdriver.Chrome()
                        driver.get('https://www.youtube.com')
                        time.sleep(2)
                        elem=driver.find_element_by_name("search_query")
                        elem.send_keys(search)
                        elem.send_keys(Keys.ENTER)
                        
                    youtubeframe=Frame(servframe,width=700,height=658,bg='#000000')
                    youtubeframe.place(x=700,y=0)
                    youtubelabel_img = PhotoImage(file = f"Frame/home_img/youtubelabel.png")
                    label = Label(image=youtubelabel_img)
                    label.image=youtubelabel_img
                    Label(youtubeframe,image=youtubelabel_img,font=("Poppins",24),background='#FFFFFF').place(x=0,y=0,width=700,height=80)

                    Label(youtubeframe,text="Search Video",font=("Poppins",24),background='#000000',foreground="#FFFFFF").place(x=0,y=170,width=700,height=80)
                    videotxt = Entry(youtubeframe,bd = 0,bg = "#FFFFFF",font=("Poppins",15),highlightthickness = 0)
                    videotxt.place(x=180,y=230,width=350)

                    Button(youtubeframe,text="Lets Go!",borderwidth = 0,highlightthickness = 0,font=("Poppins",15),command = yt_browse,background="#FFFFFF",foreground="#000000",activebackground="#FFFFFF",relief = "flat").place(x=280,y=320,width=120)

                menu.destroy()
                servframe=Frame(w,width=1400,height=658,bg='#262626')
                servframe.place(x=0,y=42)
                glabel=Label(w,text='Automated Services',font=("Poppins",24),background='#000000',foreground='#FFFFFF').place(x=400,y=0,width=600,height=42)

                servinput=Frame(servframe,width=700,height=658,bg='#000000')
                servinput.place(x=700,y=0)
                
                #gmail button
                gmail_img = PhotoImage(file = f"Frame/home_img/gmail.png")
                label = Label(image=gmail_img)
                label.image=gmail_img
                Button(servframe,borderwidth = 0,image=gmail_img,highlightthickness = 0,font=("Poppins",15),command = gmail,background="#262626",activebackground="#262626",relief = "flat").place(x=240,y=150,width=206,height=62)

                #amazon button
                amazon_img = PhotoImage(file = f"Frame/home_img/amazon.png")
                label = Label(image=amazon_img)
                label.image=amazon_img
                Button(servframe,borderwidth = 0,image=amazon_img,highlightthickness = 0,font=("Poppins",15),command = amazon,background="#262626",activebackground="#262626",relief = "flat").place(x=240,y=250,width=206,height=62)

                #flipkart button
                flipkart_img = PhotoImage(file = f"Frame/home_img/flipkart.png")
                label = Label(image=flipkart_img)
                label.image=flipkart_img
                Button(servframe,borderwidth = 0,image=flipkart_img,highlightthickness = 0,font=("Poppins",15),command = flipkart,background="#262626",activebackground="#262626",relief = "flat").place(x=240,y=350,width=206,height=62)

                #youtube button
                youtube_img = PhotoImage(file = f"Frame/home_img/youtube.png")
                label = Label(image=youtube_img)
                label.image=youtube_img
                Button(servframe,borderwidth = 0,image=youtube_img,highlightthickness = 0,font=("Poppins",15),command = youtube,background="#262626",activebackground="#262626",relief = "flat").place(x=240,y=450,width=206,height=62)

                gmail()

            def logout():
                w.destroy()
                login.Main_window('login')

            bttn(0,80,'Home','#D2E6FB','#D2E6FB',home)
            bttn(0,150,'Manage PDF','#D2E6FB','#D2E6FB',pdf_code)
            bttn(0,220,'Manage Word','#D2E6FB','#D2E6FB',word_code)
            bttn(0,290,'Weather','#D2E6FB','#D2E6FB',excel_code)
            bttn(0,360,'Services','#D2E6FB','#D2E6FB',services_code)
            bttn(0,650,'Logout','#D2E6FB','#D2E6FB',logout)
                
            def dele():
                menu.destroy()

                
            global menubtn
            menubtn=ImageTk.PhotoImage(Image.open('Frame/home_img/close.png'))

            Button(menu,image=menubtn,command=dele,border=0,background='#D2E6FB',activebackground='#D2E6FB').place(x=5,y=10)

        default_home()

        img1=ImageTk.PhotoImage(Image.open('Frame/home_img/open.png'))
        Button(w,command=toggle_win,image=img1,border=0,bg='#000000',activebackground='#000000').place(x=5,y=5)

        app_width=1400 
        app_height=700
        screen_width = w.winfo_screenwidth()
        screen_height = w.winfo_screenheight()
        x=(screen_width / 2) - (app_width/2)
        y=(screen_height / 2) - (app_height/2)
        w.geometry(f'{app_width}x{app_height}+{int(x)}+{int(y)}')
        w.deiconify()
        w.mainloop()